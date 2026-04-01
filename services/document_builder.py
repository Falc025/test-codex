from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from models.document_data import DocumentData
from services.validator import InputValidator, ValidationError
from utils.file_utils import ensure_directory


class DocumentBuilder:
    """Construye documentos Word reemplazando marcadores de plantilla."""

    def __init__(self, validator: InputValidator | None = None) -> None:
        self.validator = validator or InputValidator()

    def build_many(
        self,
        records: list[DocumentData],
        template_zero: str | Path,
        template_positive: str | Path,
        template_negative: str | Path,
        output_dir: str | Path,
        filename_prefix: str = "documento",
    ) -> list[Path]:
        tpl_zero, tpl_pos, tpl_neg = self.validator.validate_templates(
            template_zero,
            template_positive,
            template_negative,
        )
        out_dir = ensure_directory(self.validator.validate_output_dir(output_dir))

        generated: list[Path] = []
        for index, record in enumerate(records, start=1):
            template = self._resolve_template(record.total_amount(), tpl_zero, tpl_pos, tpl_neg)
            generated_path = self._build_one(
                template_path=template,
                output_dir=out_dir,
                data=record,
                filename_prefix=filename_prefix,
                sequence=index,
            )
            generated.append(generated_path)
        return generated

    def _resolve_template(self, total: float, tpl_zero: Path, tpl_pos: Path, tpl_neg: Path) -> Path:
        if total == 0:
            return tpl_zero
        if total > 0:
            return tpl_pos
        return tpl_neg

    def _build_one(
        self,
        template_path: Path,
        output_dir: Path,
        data: DocumentData,
        filename_prefix: str,
        sequence: int,
    ) -> Path:
        try:
            document = Document(template_path)
        except Exception as exc:
            raise ValidationError(f"No se pudo abrir la plantilla Word: {exc}") from exc

        placeholders = data.to_placeholders()
        self._replace_in_document(document, placeholders)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        siged = data.value("n° siged") or data.value("n_siged") or data.value("expediente") or f"fila{data.row_number}"
        safe_siged = siged.replace("/", "-").replace(" ", "_")
        output_name = f"{filename_prefix}_{sequence:03d}_{safe_siged}_{timestamp}.docx"
        output_path = output_dir / output_name

        try:
            document.save(output_path)
        except Exception as exc:
            raise ValidationError(f"No se pudo guardar el documento generado: {exc}") from exc

        return output_path

    def _replace_in_document(self, document: Document, placeholders: dict[str, str]) -> None:
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, placeholders)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, placeholders)

    @staticmethod
    def _replace_in_paragraph(paragraph, placeholders: dict[str, str]) -> None:
        full_text = paragraph.text
        if not full_text:
            return
        for token, value in placeholders.items():
            full_text = full_text.replace(token, value)

        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""
