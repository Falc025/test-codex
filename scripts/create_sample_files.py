from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook

BASE_DIR = Path(__file__).resolve().parents[1]
TEMPLATE_ZERO = BASE_DIR / "templates" / "plantilla_total_cero.docx"
TEMPLATE_POSITIVE = BASE_DIR / "templates" / "plantilla_total_positivo.docx"
TEMPLATE_NEGATIVE = BASE_DIR / "templates" / "plantilla_total_negativo.docx"
EXCEL_PATH = BASE_DIR / "sample_input.xlsx"

HEADERS = [
    "tipo",
    "año",
    "n° siged",
    "ruc",
    "razon_social",
    "domicilio",
    "periodo",
    "resul_req",
    "fecha_not_req",
    "fecha_venci",
    "bi_det",
    "alicuota",
    "apr_det",
    "apr_decla",
    "apr_omitido",
    "int_moratorio",
    "total",
    "sector",
]

ROWS = [
    ["RD", 2026, "202500291865", "20137913250", "ANGLO AMERICAN QUELLAVECO S.A.", "Calle Bernardo Monteagudo 222", "Ene-21", "25-2026-RQ-OS/UATGC", "10/01/2026", "26/02/2021", "0,00", "0,14%", "0,00", "0,00", "0,00", "0,00", "0,00", "Minería"],
    ["RD", 2026, "202500291865", "20137913250", "ANGLO AMERICAN QUELLAVECO S.A.", "Calle Bernardo Monteagudo 222", "Ago-21", "25-2026-RQ-OS/UATGC", "10/01/2026", "30/09/2021", "200000,00", "0,14%", "280,00", "10,00", "270,00", "20,00", "290,00", "Minería"],
    ["RD", 2026, "202500291865", "20137913250", "ANGLO AMERICAN QUELLAVECO S.A.", "Calle Bernardo Monteagudo 222", "Nov-21", "25-2026-RQ-OS/UATGC", "10/01/2026", "31/12/2021", "10000,00", "0,14%", "14,00", "100,00", "-86,00", "0,00", "-86,00", "Minería"],
]


def create_excel() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Datos"
    ws.append(HEADERS)
    for row in ROWS:
        ws.append(row)
    wb.save(EXCEL_PATH)


def create_template(path: Path, title: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("SIGED: {{n° siged}}")
    doc.add_paragraph("RUC: {{ruc}}")
    doc.add_paragraph("RAZÓN SOCIAL: {{razon_social}}")
    doc.add_paragraph("PERIODO: {{periodo}}")
    doc.add_paragraph("TOTAL: {{total}}")
    doc.add_paragraph("SECTOR: {{sector}}")
    doc.save(path)


def main() -> None:
    TEMPLATE_ZERO.parent.mkdir(parents=True, exist_ok=True)
    create_excel()
    create_template(TEMPLATE_ZERO, "Plantilla para TOTAL = 0")
    create_template(TEMPLATE_POSITIVE, "Plantilla para TOTAL > 0")
    create_template(TEMPLATE_NEGATIVE, "Plantilla para TOTAL < 0")
    print(f"Excel de ejemplo creado en: {EXCEL_PATH}")
    print(f"Plantillas creadas en: {TEMPLATE_ZERO.parent}")


if __name__ == "__main__":
    main()
