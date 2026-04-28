from __future__ import annotations

import re
from xml.sax.saxutils import escape as xml_escape
from dataclasses import dataclass
from pathlib import Path
from typing import Any
import xml.etree.ElementTree as ET
from zipfile import ZIP_DEFLATED, BadZipFile, ZipFile

from docx import Document

from app.utils.text_utils import extract_placeholders


@dataclass
class TemplateScan:
    placeholders: set[str]


class TemplateEngine:
    _token_pattern = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
    _word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def scan_placeholders(self, docx_path: str | Path) -> TemplateScan:
        doc = Document(docx_path)
        found: set[str] = set()
        for paragraph in self._iter_all_paragraphs(doc):
            found.update(extract_placeholders(paragraph.text))

        # También escanear placeholders en footnotes reales.
        with ZipFile(docx_path, "r") as zf:
            if "word/footnotes.xml" in zf.namelist():
                found.update(self._extract_placeholders_from_xml_part(zf.read("word/footnotes.xml")))

        return TemplateScan(placeholders=found)

    def render(self, template_path: str | Path, data: dict[str, str], output_path: str | Path) -> None:
        doc = Document(template_path)
        self.replace_placeholders_in_document(doc, data)
        doc.save(output_path)

        # Etapa específica y segura para notas al pie reales (word/footnotes.xml)
        self.replace_placeholders_in_footnotes_xml_preserving_raw(output_path, data)

        ok, errors = self.validate_docx_integrity(output_path)
        if not ok:
            raise ValueError("DOCX inválido tras procesar notas al pie: " + " | ".join(errors))

    def replace_placeholders_in_document(self, doc: Document, values: dict[str, str]) -> None:
        for container in self._iter_all_containers(doc):
            self.replace_placeholders_in_container(container, values)

    def replace_placeholders_in_container(self, container: Any, values: dict[str, str]) -> None:
        for paragraph in container.paragraphs:
            self._replace_in_paragraph(paragraph, values)
        for table in container.tables:
            self._replace_in_table(table, values)

    def _replace_in_table(self, table: Any, values: dict[str, str]) -> None:
        for row in table.rows:
            for cell in row.cells:
                self.replace_placeholders_in_container(cell, values)

    def _iter_all_containers(self, doc: Document):
        yield doc
        visited: set[int] = set()
        for section in doc.sections:
            containers = [
                section.header,
                section.footer,
                section.first_page_header,
                section.first_page_footer,
                section.even_page_header,
                section.even_page_footer,
            ]
            for container in containers:
                if id(container) in visited:
                    continue
                visited.add(id(container))
                yield container

    def _iter_all_paragraphs(self, doc: Document):
        for container in self._iter_all_containers(doc):
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p

    def _replace_in_paragraph(self, paragraph: Any, data: dict[str, str]) -> None:
        if not paragraph.runs:
            return
        original_text = "".join(run.text for run in paragraph.runs)
        replaced_text = self._token_pattern.sub(lambda m: data.get(m.group(1).strip(), ""), original_text)
        if replaced_text == original_text:
            return

        lengths = [len(run.text) for run in paragraph.runs]
        cursor = 0
        for i, run in enumerate(paragraph.runs):
            take = lengths[i]
            segment = replaced_text[cursor: cursor + take]
            run.text = segment
            cursor += len(segment)
        if cursor < len(replaced_text):
            paragraph.runs[-1].text += replaced_text[cursor:]

    def replace_placeholders_in_footnotes_xml_preserving_raw(self, docx_path: str | Path, values: dict[str, str]) -> None:
        src = Path(docx_path)
        temp = src.with_suffix(".tmp.docx")

        with ZipFile(src, "r") as zin:
            names = zin.namelist()
            footnotes_exists = "word/footnotes.xml" in names
            original_entries = {name: zin.read(name) for name in names}

        if not footnotes_exists:
            return

        footnotes_xml = original_entries["word/footnotes.xml"].decode("utf-8")
        updated_footnotes = self._replace_placeholders_in_footnotes_raw_xml(footnotes_xml, values).encode("utf-8")

        with ZipFile(temp, "w", compression=ZIP_DEFLATED) as zout:
            for name, payload in original_entries.items():
                if name == "word/footnotes.xml":
                    zout.writestr(name, updated_footnotes)
                else:
                    zout.writestr(name, payload)

        temp.replace(src)

    def _replace_placeholders_in_footnotes_raw_xml(self, xml_text: str, values: dict[str, str]) -> str:
        updated = xml_text
        for key, raw_value in values.items():
            value = xml_escape(str(raw_value if raw_value is not None else ""))
            # Caso 1: placeholder continuo.
            simple_pattern = re.compile(rf"\{{\{{\s*{re.escape(key)}\s*\}}\}}")
            updated = simple_pattern.sub(value, updated)
            # Caso 2: placeholder fragmentado entre <w:t> y posibles tags intermedios (proofErr, etc).
            updated = self.replace_fragmented_placeholder(updated, key, value)
        return updated

    def replace_fragmented_placeholder(self, xml_text: str, key: str, value: str) -> str:
        pattern = re.compile(
            rf"<w:t(?P<t1>[^>]*)>\s*\{{\{{\s*</w:t>"
            rf"(?P<m1>(?:\s|<[^>]+>)*?)"
            rf"<w:t(?P<t2>[^>]*)>\s*{re.escape(key)}\s*</w:t>"
            rf"(?P<m2>(?:\s|<[^>]+>)*?)"
            rf"<w:t(?P<t3>[^>]*)>\s*\}}\}}\s*</w:t>",
            re.DOTALL,
        )

        def repl(match: re.Match) -> str:
            attrs = match.group("t1") or ""
            return f"<w:t{attrs}>{value}</w:t>"

        return pattern.sub(repl, xml_text)

    def find_remaining_placeholders(self, docx_path: str | Path) -> dict[str, list[str]]:
        found: dict[str, list[str]] = {}
        target_parts = {"word/document.xml", "word/footnotes.xml", "word/endnotes.xml"}
        with ZipFile(docx_path, "r") as zf:
            for name in zf.namelist():
                if not (name.startswith("word/") and name.endswith(".xml")):
                    continue
                if name not in target_parts and not name.startswith("word/header") and not name.startswith("word/footer"):
                    continue
                xml_text = zf.read(name).decode("utf-8", errors="ignore")
                placeholders = sorted({m.group(1).strip() for m in self._token_pattern.finditer(xml_text)})
                if placeholders:
                    found[name] = [f"{{{{{p}}}}}" for p in placeholders]
        return found

    def validate_docx_integrity(self, docx_path: str | Path) -> tuple[bool, list[str]]:
        errors: list[str] = []
        required_parseable = ["word/document.xml", "word/footnotes.xml"]
        critical_duplicates = {"[Content_Types].xml", "word/document.xml", "word/footnotes.xml"}

        try:
            with ZipFile(docx_path, "r") as zf:
                infos = zf.infolist()
                names = [i.filename for i in infos]

                # duplicados críticos
                for item in critical_duplicates:
                    if names.count(item) > 1:
                        errors.append(f"Entrada ZIP duplicada crítica: {item}")

                # parse XML principal / footnotes si existe
                for name in required_parseable:
                    if name in names:
                        try:
                            ET.fromstring(zf.read(name))
                        except ET.ParseError as exc:
                            errors.append(f"XML inválido {name}: {exc}")

                # no placeholders pendientes en footnotes
                if "word/footnotes.xml" in names:
                    foot_raw = zf.read("word/footnotes.xml").decode("utf-8", errors="ignore")
                    remaining_foot = {m.group(1).strip() for m in self._token_pattern.finditer(foot_raw)}
                    if remaining_foot:
                        errors.append(
                            "Placeholders remanentes en word/footnotes.xml: "
                            + ", ".join(f"{{{{{k}}}}}" for k in sorted(remaining_foot))
                        )
                    if "xmlns:w=" in foot_raw and "ns0:" in foot_raw:
                        errors.append("word/footnotes.xml contiene prefijos ns0 inesperados")
        except BadZipFile as exc:
            errors.append(f"DOCX no es ZIP válido: {exc}")

        return (len(errors) == 0, errors)

    def _extract_placeholders_from_xml_part(self, xml_bytes: bytes) -> set[str]:
        try:
            root = ET.fromstring(xml_bytes)
            text_xpath = f".//{{{self._word_ns}}}t"
            text = "".join(node.text or "" for node in root.findall(text_xpath))
        except ET.ParseError:
            text = xml_bytes.decode("utf-8", errors="ignore")
        return {m.group(1).strip() for m in self._token_pattern.finditer(text)}
